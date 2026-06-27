from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_proposal():
    doc = Document()
    
    # Title
    title = doc.add_heading('DRAFT KONTEN PROPOSAL RIIM STARTUP BRIN', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Catatan: Ini adalah draf isi (teks) proposal berdasarkan ide project KOL Scouting Web App. Silakan copy-paste teks ini ke dalam template asli (68a817ce054d7.docx) agar format tabel dan lembar pengesahan tetap rapi.\n')

    # Identitas
    doc.add_heading('IDENTITAS PROPOSAL', level=1)
    doc.add_paragraph('JUDUL PROPOSAL: KOL Scout: Platform Manajemen, Kurasi, dan Analitik Influencer Berbasis Data untuk Pemberdayaan UMKM')
    doc.add_paragraph('BIDANG FOKUS: Teknologi Informasi dan Komunikasi / Ekonomi Digital')
    doc.add_paragraph('PRODUK: Platform Web App KOL Scout (KOL Database & Campaign Management)')
    
    # Ringkasan Eksekutif
    doc.add_heading('RINGKASAN EKSEKUTIF / ABSTRAK', level=1)
    doc.add_paragraph('KOL Scout adalah inovasi platform berbasis web yang dirancang untuk mendisrupsi dan mengefisiensikan cara UMKM dan Brand dalam menemukan, mengelola, dan melacak Key Opinion Leader (KOL) atau influencer. Saat ini, proses pencarian KOL lokal (seperti di area Malang Raya) dilakukan secara manual, memakan waktu lama, dan tidak didasarkan pada metrik performa yang transparan. KOL Scout hadir dengan teknologi automasi scraping data media sosial (menggunakan API), sistem filter cerdas (algoritma penyaringan akun bisnis, nano-influencer, demografi lokasi), dan Campaign Tracker (Papan Laporan Live) yang memungkinkan brand mengatur storyline, review draft, hingga melacak status tayang kampanye dalam satu pintu. Melalui pendanaan RIIM Startup, produk ini akan dikomersialisasikan sebagai platform Software as a Service (SaaS) B2B yang bertujuan mendemokratisasi akses pemasaran digital bagi UMKM dan meningkatkan ekosistem ekonomi kreatif kreator lokal.')

    # BAB 1
    doc.add_heading('BAB 1. PENDAHULUAN', level=1)
    doc.add_heading('Latar Belakang', level=2)
    doc.add_paragraph('Perkembangan ekonomi digital membuat pemasaran melalui Key Opinion Leader (KOL) menjadi salah satu strategi paling efektif bagi UMKM dan brand skala besar. Namun, proses scouting (pencarian) dan pengelolaan kampanye influencer saat ini masih sangat konvensional. Pemilik bisnis harus mencari secara manual melalui hashtag Instagram, menghadapi ketidakpastian metrik (seperti fake followers atau low engagement rate), serta mengelola ratusan data menggunakan spreadsheet tradisional. Hal ini menyebabkan inefisiensi biaya dan waktu yang signifikan, serta tingginya risiko kegagalan kampanye akibat salah pilih KOL. KOL Scout diusulkan untuk menyelesaikan *pain point* ini dengan menyediakan database instan berbasis analitik dan dasbor manajemen kampanye terpusat.')
    
    doc.add_heading('Tujuan Kegiatan', level=2)
    doc.add_paragraph('1. Jangka Pendek: Menyempurnakan sistem smart-filtering lokasi dan metrik engagement rate untuk KOL lokal secara real-time, serta melakukan validasi pasar di wilayah Malang Raya.\n'
                      '2. Jangka Menengah: Mengintegrasikan sistem Portal KOL mandiri di mana influencer dapat mendaftar dan memberikan draf konten secara langsung ke dalam platform, serta memperluas database ke tingkat nasional.\n'
                      '3. Jangka Panjang: Mengembangkan AI untuk memprediksi keberhasilan suatu kampanye berdasarkan data historis KOL, serta menjalin kerja sama strategis dengan asosiasi UMKM Nasional.')
    
    doc.add_heading('Manfaat Kegiatan', level=2)
    doc.add_paragraph('- Bagi Startup: Memperoleh validasi pasar, akselerasi pengembangan produk teknologi, dan pendapatan dari skema langganan (subscription).\n'
                      '- Bagi UMKM & Brand: Menghemat lebih dari 70% waktu pencarian KOL, menekan anggaran pemasaran yang terbuang karena salah target, dan mempermudah pemantauan ROI kampanye.\n'
                      '- Bagi Masyarakat/Kreator Lokal: Membantu micro dan nano-influencer di daerah agar lebih mudah ditemukan oleh brand dan mendapatkan peluang ekonomi.')

    # BAB 2
    doc.add_heading('BAB 2. PROFIL STARTUP', level=1)
    doc.add_heading('Visi dan Misi Startup', level=2)
    doc.add_paragraph('Visi: Menjadi infrastruktur teknologi utama yang menjembatani pertumbuhan ekonomi UMKM dengan ekosistem kreator digital di Asia Tenggara.\n'
                      'Misi: \n'
                      '1. Menyediakan data analitik KOL yang transparan, akurat, dan real-time.\n'
                      '2. Menghadirkan alat manajemen kampanye yang intuitif bagi agency dan UMKM.\n'
                      '3. Mendukung pertumbuhan nano dan micro-influencer lokal melalui visibilitas platform.')
    
    doc.add_heading('Kondisi yang diharapkan sebagai Startup berbasis riset', level=2)
    doc.add_paragraph('KOL Scout bukan sekadar direktori, melainkan platform berbasis riset big data dan algoritma penyaringan. Kami terus melakukan riset terkait natural language processing (NLP) untuk menganalisis sentimen audiens pada kolom komentar KOL dan identifikasi lokasi geografis (Location Intelligence) berdasarkan pola teks pada caption dan bio, seperti yang sudah sukses diterapkan pada filter area lokal.')

    # BAB 3
    doc.add_heading('BAB 3. PRODUK/JASA HASIL RISET', level=1)
    doc.add_paragraph('Produk utama kami adalah platform Web App KOL Scout. \n'
                      '- Keterbaharuan (Novelty): Dibandingkan dengan marketplace influencer biasa, KOL Scout memiliki algoritma "Smart Location & Niche Filter" yang secara dinamis membuang akun bisnis/merchant dan secara akurat mendeteksi lokasi KOL dari metadata teks yang tidak terstruktur.\n'
                      '- Keunggulan: Fitur Workspace Campaign yang memungkinkan kolaborasi brand dan agency dalam memantau tautan draf video, memberikan feedback revisi (Drafting, Ready for Review, Approved), dan fitur Upload Tracker untuk memantau URL live postingan KOL.\n'
                      '- Kekayaan Intelektual: Source code algoritma filtering dan arsitektur database kampanye (dalam proses pendaftaran hak cipta/paten perangkat lunak).')

    # BAB 4
    doc.add_heading('BAB 4. RENCANA BISNIS', level=1)
    doc.add_paragraph('1. Business Model: Model bisnis B2B Software as a Service (SaaS). Kami menerapkan sistem Freemium dan Subscription bulanan/tahunan (Tiered Pricing) untuk Brand/Agency berdasarkan jumlah kapasitas list KOL yang dikelola dan kedalaman analitik data.\n'
                      '2. Target Pasar: Agency pemasaran digital, UMKM yang sedang berkembang, dan perusahaan FMCG skala menengah.\n'
                      '3. Rencana Pemasaran: Melakukan edukasi pasar melalui webinar "Digital Marketing & Influencer Strategy", menawarkan free-trial untuk Agency lokal, serta membangun kemitraan dengan instansi pemerintah (Dinas Koperasi & UMKM) untuk mempercepat adopsi platform di kalangan pengusaha daerah.')

    # BAB 5
    doc.add_heading('BAB 5. RENCANA HASIL YANG AKAN DICAPAI', level=1)
    doc.add_paragraph('Tahun I: \n'
                      '- Penyelesaian pengembangan fitur KOL Portal (Influencer Login).\n'
                      '- Akuisisi 50 B2B Client berbayar (Agency/Brand) dan database 5.000 KOL tervalidasi.\n'
                      '- Pendaftaran Hak Cipta Perangkat Lunak.\n'
                      '- Integrasi API resmi dari platform media sosial utama.\n\n'
                      'Tahun II: \n'
                      '- Skalabilitas ke 5 kota besar di Indonesia.\n'
                      '- Pengembangan Modul AI Predictive Campaign ROI.\n'
                      '- Pencapaian profitabilitas operasional berkelanjutan.')

    # BAB 6
    doc.add_heading('BAB 6. PENUTUP', level=1)
    doc.add_paragraph('Pengembangan KOL Scout melalui pendanaan RIIM Startup diharapkan dapat menjadi katalis dalam mendigitalisasi sektor pemasaran untuk UMKM di Indonesia. Dengan mengatasi ketidakefisienan dalam manajemen influencer, produk inovasi kami tidak hanya memiliki potensi komersial yang tinggi tetapi juga dampak sosial-ekonomi yang nyata bagi kreator konten dan pemilik bisnis lokal.')
    
    doc.save(r'D:\Django Project\KOL_Scouting_Project\proposal\Draft_Isi_Proposal_KOL_Scout.docx')

if __name__ == '__main__':
    create_proposal()
